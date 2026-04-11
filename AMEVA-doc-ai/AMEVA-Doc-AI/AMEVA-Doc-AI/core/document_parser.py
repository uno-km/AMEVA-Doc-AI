import os, zipfile, re, zlib
import xml.etree.ElementTree as ET
import docx, openpyxl, pptx, olefile

class DocumentParser:
    @staticmethod
    def extract_all_text(path):
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == '.hwpx': return DocumentParser._extract_hwpx(path)
            elif ext == '.hwp': return DocumentParser._extract_hwp(path)
            elif ext == '.docx': return DocumentParser._extract_docx(path)
            elif ext == '.xlsx': return DocumentParser._extract_excel(path)
            elif ext == '.pptx': return DocumentParser._extract_ppt(path)
            else: return "지원하지 않는 확장자입니다."
        except Exception as e:
            return f"문서 추출 실패: {str(e)}"

    @staticmethod
    def _extract_hwpx(path):
        texts = []
        with zipfile.ZipFile(path, 'r') as z:
            sections = sorted([f for f in z.namelist() if 'Contents/section' in f])
            for s in sections:
                with z.open(s) as f:
                    for t in ET.parse(f).getroot().iter():
                        if t.tag.endswith('t') and t.text: texts.append(t.text)
        return " ".join(texts)

    @staticmethod
    def _extract_hwp(path):
        f = olefile.OleFileIO(path)
        dirs = f.listdir()
        text = ""
        for d in dirs:
            if d[0] == 'BodyText':
                stream = f.openstream(d)
                data = stream.read()
                try:
                    decompressed = zlib.decompress(data, -15)
                    decoded = decompressed.decode('utf-16le', errors='ignore')
                    clean = re.sub(r'[^\uAC00-\uD7A3a-zA-Z0-9\s\.\,\?\!\-\(\)\[\]\"\'\%\:\/\~\=]', '', decoded)
                    text += clean + " "
                except: pass
        return text

    @staticmethod
    def _extract_docx(path):
        doc = docx.Document(path)
        texts = []
        for p in doc.paragraphs:
            if p.text.strip(): texts.append(p.text.strip())
            
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                seen_cells = set()
                for cell in row.cells:
                    if cell in seen_cells:
                        row_data.append(" ")
                        continue
                    seen_cells.add(cell)
                    clean_text = cell.text.replace('|', '/').strip()
                    clean_text = re.sub(r'\s+', ' ', clean_text)
                    row_data.append(clean_text if clean_text else " ")
                if any(cell.strip() for cell in row_data):
                    texts.append("| " + " | ".join(row_data) + " |")
            texts.append("")
        return "\n".join(texts)

    @staticmethod
    def _extract_excel(path):
        wb = openpyxl.load_workbook(path, data_only=True)
        texts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) for cell in row if cell is not None]
                if row_data: texts.append(" | ".join(row_data))
        return "\n".join(texts)

    @staticmethod
    def _extract_ppt(path):
        prs = pptx.Presentation(path)
        texts = []
        for slide in prs.slides:
            for shape in slide.shapes:
                if hasattr(shape, "has_text_frame") and shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        if paragraph.text.strip():
                            texts.append(paragraph.text.strip())
        return "\n".join(texts)
