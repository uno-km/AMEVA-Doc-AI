import os, zipfile, re, zlib
import xml.etree.ElementTree as ET
import docx, openpyxl, pptx, olefile

class DocumentParser:
    @staticmethod
    def extract_all_text(path):
        ext = os.path.splitext(path)[1].lower()
        try:
            if ext == '.hwpx': return DocumentParser._extract_hwpx(path)
            elif ext == '.hwp': return DocumentParser._extract_hwp(path)
            elif ext == '.docx': return DocumentParser._extract_docx(path)
            elif ext == '.xlsx': return DocumentParser._extract_excel(path)
            elif ext == '.pptx': return DocumentParser._extract_ppt(path)
            else: return "지원하지 않는 확장자입니다."
        except Exception as e:
            return f"문서 추출 실패: {str(e)}"

    @staticmethod
    def _extract_hwpx(path):
        texts = []
        with zipfile.ZipFile(path, 'r') as z:
            sections = sorted([f for f in z.namelist() if 'Contents/section' in f])
            for s in sections:
                with z.open(s) as f:
                    for t in ET.parse(f).getroot().iter():
                        if t.tag.endswith('t') and t.text: texts.append(t.text)
        return " ".join(texts)

    @staticmethod
    def _extract_hwp(path):
        f = olefile.OleFileIO(path)
        dirs = f.listdir()
        text = ""
        for d in dirs:
            if d[0] == 'BodyText':
                stream = f.openstream(d)
                data = stream.read()
                try:
                    decompressed = zlib.decompress(data, -15)
                    decoded = decompressed.decode('utf-16le', errors='ignore')
                    clean = re.sub(r'[^\uAC00-\uD7A3a-zA-Z0-9\s\.\,\?\!\-\(\)\[\]\"\'\%\:\/\~\=]', '', decoded)
                    text += clean + " "
                except: pass
        return text

    @staticmethod
    def _extract_docx(path):
        doc = docx.Document(path)
        texts = []
        
        # 1. 일반 문단 추출
        for p in doc.paragraphs:
            if p.text.strip():
                texts.append(p.text.strip())
                
        # 2. 표(Table) 완벽 평탄화 로직
        for table in doc.tables:
            for row in table.rows:
                row_data = []
                seen_cells = set() # 병합된 셀 중복 추출 방지용 메모리
                
                for cell in row.cells:
                    if cell in seen_cells:
                        # 이미 읽은 병합 셀이라면, 마크다운 표 구조 유지를 위해 빈칸만 삽입
                        row_data.append(" ") 
                        continue
                        
                    seen_cells.add(cell)

                    # python-docx의 cell.text는 내부에 숨겨진 '중첩 표'의 텍스트까지 전부 긁어옵니다.
                    # 마크다운 표가 깨지지 않도록 내부의 파이프(|)를 슬래시(/)로 바꾸고, 
                    # 모든 종류의 줄바꿈(\n)을 띄어쓰기 한 칸으로 강제 압축(평탄화)합니다.
                    clean_text = cell.text.replace('|', '/').strip()
                    clean_text = re.sub(r'\s+', ' ', clean_text) 
                    
                    row_data.append(clean_text if clean_text else " ")
                
                # 줄의 모든 칸이 비어있지 않은 경우에만 표에 추가
                if any(cell.strip() for cell in row_data):
                    texts.append("| " + " | ".join(row_data) + " |")
            
            texts.append("") # 표 하나가 끝나면 줄바꿈 추가하여 다음 문단과 분리
            
        return "\n".join(texts)

    @staticmethod
    def _extract_excel(path):
        wb = openpyxl.load_workbook(path, data_only=True)
        texts = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(values_only=True):
                row_data = [str(cell) for cell in row if cell is not None]
                if row_data: texts.append(" | ".join(row_data))
        return "\n".join(texts)

    @staticmethod
    def _extract_ppt(path):
        prs = pptx.Presentation(path)
        texts = []
        for slide in prs.slides:
            for shape in slide.shapes:
                if hasattr(shape, "has_text_frame") and shape.has_text_frame:
                    for paragraph in shape.text_frame.paragraphs:
                        if paragraph.text.strip():
                            texts.append(paragraph.text.strip())
        return "\n".join(texts)